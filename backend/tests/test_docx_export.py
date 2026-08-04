"""Conversion Markdown -> Word du rapport composite (§app/reports/docx_export.py) : même
sous-ensemble Markdown que frontend/src/components/Markdown.tsx sait déjà afficher à l'écran."""
from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient

from app.reports.docx_export import markdown_to_docx, report_docx_filename
from app.store.db import session_scope
from app.store.repository import create_dossier


def _doc(markdown_text: str, *, title: str | None = None) -> Document:
    return Document(io.BytesIO(markdown_to_docx(markdown_text, title=title)))


def test_headings_map_to_docx_heading_levels():
    doc = _doc("# Titre 1\n## Titre 2\n### Titre 3")
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert styles == ["Heading 1", "Heading 2", "Heading 3"]


def test_consecutive_headings_without_blank_lines_are_not_merged():
    """Le LLM enchaîne parfois plusieurs titres sans ligne vide entre eux — un découpage par bloc
    les fusionnerait en un seul paragraphe (§docstring du module, même raison que Markdown.tsx)."""
    doc = _doc("### Objectifs\n#### Environnementaux\ntexte du paragraphe")
    texts = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]
    assert texts == [
        ("Heading 3", "Objectifs"),
        ("Heading 4", "Environnementaux"),
        ("Normal", "texte du paragraphe"),
    ]


def test_bold_inline_is_split_into_a_bold_run():
    doc = _doc("Ceci est **important** à retenir.")
    para = next(p for p in doc.paragraphs if p.text.strip())
    bold_runs = [r for r in para.runs if r.bold]
    assert any(r.text == "important" for r in bold_runs)
    assert not any(r.bold for r in para.runs if r.text != "important")


def test_gfm_table_becomes_a_docx_table_with_bold_header():
    md = "| Donnée | Valeur |\n|---|---|\n| Montant | 12 M€ |\n| Délai | 24 mois |"
    doc = _doc(md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["Donnée", "Valeur"]
    assert [c.text for c in table.rows[1].cells] == ["Montant", "12 M€"]
    assert [c.text for c in table.rows[2].cells] == ["Délai", "24 mois"]
    header_runs = table.rows[0].cells[0].paragraphs[0].runs
    assert all(r.bold for r in header_runs)


def test_bulleted_list_with_sub_bullets():
    md = "- Premier point\n  - Sous-point A\n  - Sous-point B\n- Deuxième point"
    doc = _doc(md)
    items = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]
    assert items == [
        ("List Bullet", "Premier point"),
        ("List Bullet 2", "Sous-point A"),
        ("List Bullet 2", "Sous-point B"),
        ("List Bullet", "Deuxième point"),
    ]


def test_numbered_list_uses_list_number_style():
    doc = _doc("1. Un\n2. Deux")
    items = [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]
    assert items == [("List Number", "Un"), ("List Number", "Deux")]


def test_paragraph_lines_are_joined_with_a_break_not_merged_into_one_line():
    doc = _doc("Ligne un\nLigne deux")
    para = next(p for p in doc.paragraphs if p.text.strip())
    assert para.text == "Ligne un\nLigne deux"


def test_optional_title_is_added_as_a_word_title_style():
    doc = _doc("# Section", title="Rapport d'analyse")
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.paragraphs[0].text == "Rapport d'analyse"


def test_filename_is_derived_from_the_original_dossier_name():
    class _FakeDossier:
        id = "abc123"
        original_filename = "DCE Marly / conservatoire.zip"

    assert report_docx_filename(_FakeDossier()) == "rapport_DCE Marly _ conservatoire.docx"


def test_export_endpoint_serves_a_docx(isolated_workspace):
    from app.main import app

    with session_scope() as s:
        dossier = create_dossier(s, original_filename="DCE Marly.zip")
        dossier_id = dossier.id

    client = TestClient(app)
    response = client.post(
        f"/api/dossiers/{dossier_id}/report/export.docx",
        json={"markdown": "# Rapport\n\nContenu."},
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in response.headers["content-disposition"]
    doc = Document(io.BytesIO(response.content))
    assert doc.paragraphs[0].text == "Rapport"


def test_export_endpoint_404s_on_unknown_dossier(isolated_workspace):
    from app.main import app

    client = TestClient(app)
    response = client.post("/api/dossiers/inconnu/report/export.docx", json={"markdown": "# X"})
    assert response.status_code == 404

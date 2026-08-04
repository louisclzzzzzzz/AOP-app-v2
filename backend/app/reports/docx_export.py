"""Convertit en document Word (.docx) le sous-ensemble Markdown produit par les rapports IA et
assemblé côté frontend (arborescence, pièces, tableau d'extraction, synthèse projet, audit des
risques — §frontend/src/components/ExtractionSheet.tsx `handleDownloadReport`) : mêmes
constructions que celles que `frontend/src/components/Markdown.tsx` sait déjà afficher à l'écran
(titres `#` à `######`, tableaux GFM, listes à puces/numérotées avec un niveau de sous-puces
indentées, gras `**...**`, paragraphes) — pas un convertisseur Markdown généraliste, et
volontairement tenu au plus près de ce parseur (mêmes regex, même logique ligne par ligne) pour
que le .docx reflète fidèlement ce que l'utilisateur a déjà vu à l'écran.

Aucune dépendance à `pandoc` (absent de l'image Docker — l'ajouter demanderait une couche
apt-get) : `python-docx` suffit pour ce sous-ensemble borné, déjà une dépendance du projet
(lecture des .docx à l'ingestion, app/ingestion/text_extraction.py).

Ne fait AUCUNE requête base de données : le Markdown à convertir est entièrement fourni par
l'appelant (assemblé côté frontend depuis l'état déjà chargé) — ce module ne fait que le rendu,
sans dupliquer la logique d'assemblage du rapport (arborescence, tableaux) côté serveur."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.text.paragraph import Paragraph

from app.store.models import Dossier

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_TABLE_SEPARATOR_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")
_UL_RE = re.compile(r"^[-*]\s+(.*)$")
_OL_RE = re.compile(r"^\d+\.\s+(.*)$")
_INDENTED_RE = re.compile(r"^\s+\S")
_BOLD_SPLIT_RE = re.compile(r"(\*\*[^*]+\*\*)")


def _add_inline(paragraph: Paragraph, text: str) -> None:
    """Ajoute `text` au paragraphe sous forme de runs, gérant le seul style inline supporté
    (`**gras**`) — miroir de `parseInline` (Markdown.tsx)."""
    for part in _BOLD_SPLIT_RE.split(text):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        if is_bold:
            run.bold = True


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def markdown_to_docx(markdown_text: str, *, title: str | None = None) -> bytes:
    """Rend `markdown_text` en un document Word, retourné en bytes (buffer mémoire, comme
    `build_extraction_workbook` pour l'export Excel — pas de fichier temporaire sur disque)."""
    document = Document()
    if title:
        document.add_heading(title, level=0)

    lines = [line.rstrip() for line in markdown_text.replace("\r\n", "\n").split("\n")]
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            p = document.add_heading("", level=level)
            _add_inline(p, heading.group(2))
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < n and _TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
            header = _split_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1

            table = document.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for cell, text in zip(table.rows[0].cells, header):
                _add_inline(cell.paragraphs[0], text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row_values in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row_values):
                    _add_inline(cell.paragraphs[0], text)
            document.add_paragraph()
            continue

        ul_match = _UL_RE.match(line)
        ol_match = _OL_RE.match(line)
        if ul_match or ol_match:
            ordered = ol_match is not None
            while i < n:
                marker = _OL_RE.match(lines[i]) if ordered else _UL_RE.match(lines[i])
                if not marker:
                    break
                p = document.add_paragraph(style="List Number" if ordered else "List Bullet")
                _add_inline(p, marker.group(1))
                i += 1
                while i < n and _INDENTED_RE.match(lines[i]):
                    continuation = lines[i].strip()
                    sub_ul = _UL_RE.match(continuation)
                    if sub_ul:
                        sub_p = document.add_paragraph(style="List Bullet 2")
                        _add_inline(sub_p, sub_ul.group(1))
                    else:
                        # Suite du même item (texte continu, pas une sous-puce) : ajoutée au même
                        # paragraphe, séparée par un saut de ligne — miroir des `extraLines`
                        # (`<br/>`) de Markdown.tsx.
                        p.add_run().add_break()
                        _add_inline(p, continuation)
                    i += 1
            continue

        # Paragraphe : lignes consécutives non vides, non titre/liste/tableau — chacune sur sa
        # propre ligne au sein du même paragraphe (comme les `<br/>` de Markdown.tsx).
        para_lines: list[str] = []
        while (
            i < n
            and lines[i].strip()
            and not _HEADING_RE.match(lines[i])
            and not _UL_RE.match(lines[i])
            and not _OL_RE.match(lines[i])
            and not lines[i].strip().startswith("|")
        ):
            para_lines.append(lines[i].strip())
            i += 1
        p = document.add_paragraph()
        for li, text in enumerate(para_lines):
            if li > 0:
                p.add_run().add_break()
            _add_inline(p, text)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def report_docx_filename(dossier: Dossier) -> str:
    """Nom de fichier proposé au téléchargement — même schéma que
    `extraction_workbook_filename` (§app/extraction/excel_export.py)."""
    stem = (dossier.original_filename or dossier.id).rsplit(".", 1)[0]
    safe = "".join(c if c.isalnum() or c in " -_" else "_" for c in stem).strip() or dossier.id
    return f"rapport_{safe}.docx"

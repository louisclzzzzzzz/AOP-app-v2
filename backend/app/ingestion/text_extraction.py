"""Routage de l'extraction de texte par type de fichier (§3.3 du PLAN).

Règles (précision prime) :
- PDF texte natif : extraction directe via pdfplumber (rapide, exacte) + OCR de contrôle
  Mistral sur les pages à faible densité de texte (PDF scannés partiels).
- PDF dont la densité moyenne est très faible sur tout le document (scan complet / plan) :
  OCR Mistral systématique sur toutes les pages.
- Image : OCR Mistral systématique.
- DOCX / feuille de calcul native (.xlsx/.csv) : extraction directe (texte né numérique,
  confiance 1.0 — ce n'est pas une inférence).
- DOC (legacy binaire) : conversion via LibreOffice (si disponible) vers PDF, puis même
  routage que pour un PDF. Si LibreOffice indisponible : échec explicite, jamais de texte
  deviné.
"""
from __future__ import annotations

import csv
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import docx
import openpyxl
import pdfplumber
import xlrd

from app.ocr.service import get_ocr_model_version, run_ocr
from app.settings import get_models_config
from app.store.models import FileCategory, TextExtractionMethod


@dataclass
class ExtractionOutcome:
    method: str
    combined_text: str
    avg_confidence: float | None
    page_count: int | None
    char_count: int
    pages_meta: list[dict[str, Any]] = field(default_factory=list)
    model_name: str | None = None
    model_version: str | None = None
    error: str | None = None
    raw_json: str | None = None


# --- PDF (natif + OCR de contrôle / OCR complet) ----------------------------------

def _extract_native_pdf_pages(path: Path) -> list[str]:
    with pdfplumber.open(path) as pdf:
        return [(page.extract_text() or "") for page in pdf.pages]


def _full_ocr(path: Path, *, method_label: str) -> ExtractionOutcome:
    outcome = run_ocr(path, pages=None)
    pages_meta = [
        {
            "index": p.index,
            "method": "ocr",
            "confidence": p.avg_confidence,
            "char_count": p.char_count,
        }
        for p in outcome.pages
    ]
    return ExtractionOutcome(
        method=method_label,
        combined_text=outcome.combined_markdown,
        avg_confidence=outcome.avg_confidence,
        page_count=len(outcome.pages),
        char_count=len(outcome.combined_markdown),
        pages_meta=pages_meta,
        model_name=outcome.model,
        model_version=get_ocr_model_version(),
        raw_json=outcome.raw_json,
    )


def _native_only_pdf(page_texts: list[str]) -> ExtractionOutcome:
    """Mode expérimental « OCR différé » (§5 OPTIMISATION.md, phase 4) : aucun appel OCR, on
    garde le texte natif tel quel, quelle que soit sa densité. Un document sans aucun texte
    natif (scan complet, PDF chiffré) reste marqué `deferred` — à ré-extraire à la demande si un
    document s'avère concerné par l'extraction (`ensure_document_ocr`)."""
    total_chars = sum(len(t.strip()) for t in page_texts)
    if total_chars == 0:
        return ExtractionOutcome(
            method=TextExtractionMethod.DEFERRED.value,
            combined_text="",
            avg_confidence=None,
            page_count=len(page_texts) or None,
            char_count=0,
        )
    combined = "\n\n".join(f"<!-- page {i} -->\n{t}" for i, t in enumerate(page_texts))
    pages_meta = [
        {"index": i, "method": "native_pdf", "confidence": 1.0, "char_count": len(t)}
        for i, t in enumerate(page_texts)
    ]
    return ExtractionOutcome(
        method=TextExtractionMethod.NATIVE_PDF.value,
        combined_text=combined,
        avg_confidence=1.0,
        page_count=len(page_texts),
        char_count=len(combined),
        pages_meta=pages_meta,
    )


def extract_pdf(path: Path, *, allow_ocr: bool = True) -> ExtractionOutcome:
    cfg = get_models_config()["text_extraction"]
    try:
        page_texts = _extract_native_pdf_pages(path)
    except Exception:
        page_texts = []

    if not allow_ocr:
        return _native_only_pdf(page_texts)

    page_count = len(page_texts)
    if page_count == 0:
        # PDF illisible nativement (chiffré, corrompu, ou entièrement image) -> OCR intégral
        return _full_ocr(path, method_label=TextExtractionMethod.OCR.value)

    threshold_scanned = cfg["scanned_pdf_density_threshold"]
    threshold_native = cfg["native_text_density_threshold"]
    total_chars = sum(len(t.strip()) for t in page_texts)
    avg_density = total_chars / page_count

    if avg_density < threshold_scanned:
        return _full_ocr(path, method_label=TextExtractionMethod.OCR.value)

    low_density_pages = [i for i, t in enumerate(page_texts) if len(t.strip()) < threshold_native]

    if not low_density_pages:
        pages_meta = [
            {"index": i, "method": "native_pdf", "confidence": 1.0, "char_count": len(t)}
            for i, t in enumerate(page_texts)
        ]
        combined = "\n\n".join(f"<!-- page {i} -->\n{t}" for i, t in enumerate(page_texts))
        return ExtractionOutcome(
            method=TextExtractionMethod.NATIVE_PDF.value,
            combined_text=combined,
            avg_confidence=1.0,
            page_count=page_count,
            char_count=len(combined),
            pages_meta=pages_meta,
        )

    # Mixte : OCR de contrôle uniquement sur les pages à faible densité de texte natif
    ocr_outcome = run_ocr(path, pages=low_density_pages)
    ocr_by_index = {p.index: p for p in ocr_outcome.pages}

    merged_texts: list[str] = []
    pages_meta = []
    for i, native_text in enumerate(page_texts):
        if i in ocr_by_index:
            p = ocr_by_index[i]
            merged_texts.append(p.markdown)
            pages_meta.append(
                {"index": i, "method": "ocr", "confidence": p.avg_confidence, "char_count": p.char_count}
            )
        else:
            merged_texts.append(native_text)
            pages_meta.append(
                {"index": i, "method": "native_pdf", "confidence": 1.0, "char_count": len(native_text)}
            )

    combined = "\n\n".join(f"<!-- page {i} -->\n{t}" for i, t in enumerate(merged_texts))
    confidences = [m["confidence"] for m in pages_meta if m["confidence"] is not None]
    avg_confidence = sum(confidences) / len(confidences) if confidences else None

    return ExtractionOutcome(
        method=TextExtractionMethod.MIXED_PDF.value,
        combined_text=combined,
        avg_confidence=avg_confidence,
        page_count=page_count,
        char_count=len(combined),
        pages_meta=pages_meta,
        model_name=ocr_outcome.model,
        model_version=get_ocr_model_version(),
        raw_json=ocr_outcome.raw_json,
    )


# --- Image (OCR systématique) -----------------------------------------------------

def extract_image(path: Path, *, allow_ocr: bool = True) -> ExtractionOutcome:
    if not allow_ocr:
        # Une image n'a par définition aucun texte natif : sans OCR, rien à en tirer pour
        # l'instant — différé, à ré-extraire à la demande si le document devient concerné.
        return ExtractionOutcome(
            method=TextExtractionMethod.DEFERRED.value,
            combined_text="",
            avg_confidence=None,
            page_count=None,
            char_count=0,
        )
    return _full_ocr(path, method_label=TextExtractionMethod.OCR.value)


# --- DOCX (natif) ------------------------------------------------------------------

def extract_docx(path: Path) -> ExtractionOutcome:
    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    combined = "\n".join(parts)
    has_text = bool(combined.strip())
    return ExtractionOutcome(
        method=TextExtractionMethod.DOCX_NATIVE.value,
        combined_text=combined,
        avg_confidence=1.0 if has_text else None,
        page_count=None,
        char_count=len(combined),
        pages_meta=[],
        error=None if has_text else "Aucun texte extrait du .docx (probablement du contenu image) — à vérifier manuellement",
    )


# --- DOC legacy (conversion LibreOffice puis routage PDF) --------------------------

def _find_soffice() -> str | None:
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    mac_default = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if mac_default.exists():
        return str(mac_default)
    return None


def extract_doc(path: Path, *, allow_ocr: bool = True) -> ExtractionOutcome:
    soffice = _find_soffice()
    if soffice is None:
        return ExtractionOutcome(
            method=TextExtractionMethod.DOC_CONVERTED.value,
            combined_text="",
            avg_confidence=None,
            page_count=None,
            char_count=0,
            error=(
                "Conversion .doc impossible : LibreOffice (soffice) n'est pas installé sur ce "
                "poste. Installez LibreOffice pour activer la conversion automatique, ou "
                "fournissez une version .docx/.pdf de ce document."
            ),
        )
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_dir), str(path)],
                check=True,
                capture_output=True,
                timeout=180,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            return ExtractionOutcome(
                method=TextExtractionMethod.DOC_CONVERTED.value,
                combined_text="",
                avg_confidence=None,
                page_count=None,
                char_count=0,
                error=f"Échec de la conversion LibreOffice (.doc -> .pdf) : {exc}",
            )
        converted_pdf = tmp_dir / f"{path.stem}.pdf"
        if not converted_pdf.exists():
            return ExtractionOutcome(
                method=TextExtractionMethod.DOC_CONVERTED.value,
                combined_text="",
                avg_confidence=None,
                page_count=None,
                char_count=0,
                error="Conversion LibreOffice : fichier PDF de sortie introuvable",
            )
        outcome = extract_pdf(converted_pdf, allow_ocr=allow_ocr)
        if outcome.method != TextExtractionMethod.DEFERRED.value:
            outcome.method = TextExtractionMethod.DOC_CONVERTED.value
        return outcome


# --- Feuilles de calcul (natif) -----------------------------------------------------

def _extract_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"## {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xls(path: Path) -> str:
    wb = xlrd.open_workbook(str(path))
    parts: list[str] = []
    for sheet in wb.sheets():
        parts.append(f"## {sheet.name}")
        for row_idx in range(sheet.nrows):
            cells = [str(c) for c in sheet.row_values(row_idx) if str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            with open(path, encoding=encoding, newline="") as f:
                rows = list(csv.reader(f))
            return "\n".join(" | ".join(cell for cell in row if cell) for row in rows if any(row))
        except UnicodeDecodeError:
            continue
    return ""


def extract_spreadsheet(path: Path) -> ExtractionOutcome:
    ext = path.suffix.lower()
    try:
        if ext == ".csv":
            combined = _extract_csv(path)
        elif ext == ".xls":
            combined = _extract_xls(path)
        else:
            combined = _extract_xlsx(path)
    except Exception as exc:
        return ExtractionOutcome(
            method=TextExtractionMethod.SPREADSHEET_NATIVE.value,
            combined_text="",
            avg_confidence=None,
            page_count=None,
            char_count=0,
            error=f"Échec de lecture du fichier tableur : {exc}",
        )
    has_text = bool(combined.strip())
    return ExtractionOutcome(
        method=TextExtractionMethod.SPREADSHEET_NATIVE.value,
        combined_text=combined,
        avg_confidence=1.0 if has_text else None,
        page_count=None,
        char_count=len(combined),
        pages_meta=[],
        error=None if has_text else "Feuille de calcul vide ou illisible",
    )


# --- Dispatcher ----------------------------------------------------------------------

def extract_text_for_file(path: Path, category: str, *, allow_ocr: bool = True) -> ExtractionOutcome:
    if category == FileCategory.PDF.value:
        return extract_pdf(path, allow_ocr=allow_ocr)
    if category == FileCategory.IMAGE.value:
        return extract_image(path, allow_ocr=allow_ocr)
    if category == FileCategory.DOCX.value:
        return extract_docx(path)
    if category == FileCategory.DOC.value:
        return extract_doc(path, allow_ocr=allow_ocr)
    if category == FileCategory.SPREADSHEET.value:
        return extract_spreadsheet(path)
    raise ValueError(f"Catégorie non prise en charge pour extraction de texte : {category}")
